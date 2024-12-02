import re
import pandas as pd
from docx import Document

Template_Word_File= 'template.docx'
Start_Tag="{{"
End_Tag="}}"

def FindPattern(Template_Word_File,Start_Tag,End_Tag):
    
    doc = Document(Template_Word_File)
    pattern = Start_Tag+"(.*?)"+End_Tag
    matches = []

    for paragraph in doc.paragraphs:
        found = re.findall(pattern, paragraph.text)
        matches.extend(found)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                found = re.findall(pattern, cell.text)
                matches.extend(found)

    return matches

def WriteDocx(Template_Word_File, excel_file, placeholders):

    df = pd.read_excel(excel_file)

    for _, row in df.iterrows():
        doc = Document(Template_Word_File)

        placeholder_values = {
            f"{{{{{placeholder}}}}}": str(row.get(placeholder, "")).strip() for placeholder in placeholders
        }

        for paragraph in doc.paragraphs:
            runs_text = []

            for run in paragraph.runs:
                runs_text.append(run.text)

            combined_text = "".join(runs_text)

            for placeholder, new_text in placeholder_values.items():
                if placeholder in combined_text:
                    combined_text = combined_text.replace(placeholder, new_text)

                    char_idx = 0
                    for run in paragraph.runs:
                        run_length = len(run.text)
                        run.text = combined_text[char_idx:char_idx + run_length]
                        char_idx += run_length

        for table in doc.tables:
            for row in  table.rows:
                for cell in row.cells:
                    for placeholder, new_text in placeholder_values.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, new_text)

        name = placeholder_values.get("{{Name}}", "Unknown")
        output_file = f'output_{name}.docx'
        doc.save(output_file)
        print(f"Saved: {output_file}")

placeholders_found = FindPattern(Template_Word_File,Start_Tag,End_Tag)
print("Placeholders found in the Word template:")
for placeholder in placeholders_found:
    print(placeholder)

Excel_File = 'data.xlsx'
WriteDocx(Template_Word_File, Excel_File, placeholders_found)